# 文档导出服务
# 将 Markdown 文本导出为 Word（.docx）或 PDF
# 用户未指定样式时，使用 LaTeX 风格的简洁排版

import os
import uuid
import re
from datetime import datetime


def _parse_markdown_blocks(md_text: str):
    """
    将 Markdown 解析为结构化块列表
    每块是 dict: { type: 'h1'|'h2'|'h3'|'p'|'list_ol'|'list_ul'|'quote'|'code', text: str, items: [...] }
    """
    blocks = []
    lines = md_text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            blocks.append({"type": f"h{min(level, 4)}", "text": m.group(2).strip()})
            i += 1
            continue

        # 引用
        m = re.match(r"^>\s*(.*)$", line)
        if m:
            quote_lines = [m.group(1)]
            i += 1
            while i < len(lines) and re.match(r"^>\s*", lines[i]):
                quote_lines.append(re.sub(r"^>\s*", "", lines[i]))
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(quote_lines)})
            continue

        # 有序列表
        m = re.match(r"^\d+\.\s+(.*)$", line)
        if m:
            items = [m.group(1)]
            i += 1
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "list_ol", "items": items})
            continue

        # 无序列表
        m = re.match(r"^[-*]\s+(.*)$", line)
        if m:
            items = [m.group(1)]
            i += 1
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(re.sub(r"^[-*]\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "list_ul", "items": items})
            continue

        # 代码块
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            i += 1
            continue

        # 普通段落（可能跨行）
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6}\s|>\s|\d+\.\s|[-*]\s|```)", nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append({"type": "p", "text": " ".join(para_lines)})

    return blocks


def _inline_md_to_html(text: str) -> str:
    """行内 Markdown 转 HTML（用于 docx）"""
    # 转义 HTML
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # **bold**
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # *italic*
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</text>".replace("<i>", "<i>"), text)
    # `code`
    text = re.sub(r"`(.+?)`", r'<span style="font-family: Consolas, monospace; background: #f0f0f0;">\1</span>', text)
    return text


# ============== DOCX 导出 ==============

def export_to_docx(title: str, markdown_text: str, output_dir: str = None, style: str = "report") -> str:
    """
    将 Markdown 导出为 .docx 文件
    :param style: 'report'（默认，带文档标题 + 生成时间）或 'clean'（只输出内容，适合简历/翻译）
    :return: 生成的文件绝对路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if output_dir is None:
        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "static", "exports"
        )
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # 设置全局字体
    normal_style = doc.styles["Normal"]
    try:
        normal_style.font.name = "Microsoft YaHei"
    except Exception:
        pass
    normal_style.font.size = Pt(11)

    # 页面边距
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # 文档标题（report 风格才加）
    if style == "report" and title:
        title_para = doc.add_heading(level=0)
        run = title_para.add_run(title)
        run.font.color.rgb = RGBColor(0x10, 0x00, 0xA9)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub_para = doc.add_paragraph()
        sub_run = sub_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    elif style == "clean" and title:
        # clean 风格：只在顶部加一个小标题（不额外加日期等），用于简历/翻译文档自身名称
        # 注意：如果内容里已经有 # 标题，不重复添加文档级标题，避免重复
        doc.add_paragraph()

    blocks = _parse_markdown_blocks(markdown_text)

    for block in blocks:
        t = block["type"]
        if t == "h1":
            p = doc.add_heading(block["text"], level=1)
        elif t == "h2":
            p = doc.add_heading(block["text"], level=2)
        elif t == "h3":
            p = doc.add_heading(block["text"], level=3)
        elif t == "h4":
            p = doc.add_heading(block["text"], level=4)
        elif t == "p":
            html = _inline_md_to_html(block["text"])
            p = doc.add_paragraph()
            _add_runs_from_html(p, html)
        elif t == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            run = p.add_run(block["text"])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif t == "list_ol":
            for idx, item in enumerate(block["items"], 1):
                p = doc.add_paragraph(f"{idx}. {_inline_md_to_html(item)}", style="List Number")
        elif t == "list_ul":
            for item in block["items"]:
                p = doc.add_paragraph(f"• {_inline_md_to_html(item)}", style="List Bullet")
        elif t == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    filename = f"{title or 'export'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.docx"
    filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def _add_runs_from_html(paragraph, html):
    """将简单的 HTML（<b>、<i>、<span>）转换为 docx runs"""
    from docx.shared import RGBColor
    # 简化：把 <b>...</b> 切成 run
    pattern = re.compile(r"<(/?b|/?i|/?span[^>]*)>([^<]*)")
    pos = 0
    bold = False
    italic = False
    for m in pattern.finditer(html):
        # 前面文本
        if m.start() > pos:
            text = html[pos:m.start()]
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        tag = m.group(1)
        if tag == "b":
            bold = True
        elif tag == "/b":
            bold = False
        elif tag == "i":
            italic = True
        elif tag == "/i":
            italic = False
        # 跳过 span 标签中的内容
        if "span" in tag:
            run = paragraph.add_run(m.group(2))
            run.bold = bold
            run.italic = italic
        pos = m.end()
    if pos < len(html):
        run = paragraph.add_run(html[pos:])
        run.bold = bold
        run.italic = italic


# 给 _add_runs_from_html 添加为模块函数
_add_runs_from_html = _add_runs_from_html


# ============== PDF 导出 ==============

def export_to_pdf(title: str, markdown_text: str, output_dir: str = None, style: str = "report") -> str:
    """
    将 Markdown 导出为 PDF
    :param style: 'report'（默认，带文档标题 + 生成时间）或 'clean'（只输出内容）
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem

    _register_chinese_font()

    if output_dir is None:
        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "static", "exports"
        )
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{title or 'export'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.pdf"
    filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
    filepath = os.path.join(output_dir, filename)

    # 页面边距根据 style 调整
    margins = {"top": 2, "bottom": 2, "left": 2, "right": 2}
    if style == "clean":
        margins = {"top": 1.8, "bottom": 1.8, "left": 1.8, "right": 1.8}

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        leftMargin=margins["left"] * cm,
        rightMargin=margins["right"] * cm,
        topMargin=margins["top"] * cm,
        bottomMargin=margins["bottom"] * cm,
        title=title or "文档",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName="ZHFont",
        fontSize=22,
        textColor=HexColor("#1000A9"),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h1_style = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontName="ZHFont",
        fontSize=18,
        textColor=HexColor("#1000A9"),
        spaceBefore=14,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontName="ZHFont",
        fontSize=15,
        textColor=HexColor("#3B4FE0"),
        spaceBefore=10,
        spaceAfter=8,
    )
    h3_style = ParagraphStyle(
        "CustomH3",
        parent=styles["Heading3"],
        fontName="ZHFont",
        fontSize=13,
        textColor=HexColor("#4CD7F6"),
        spaceBefore=8,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontName="ZHFont",
        fontSize=11,
        leading=16,
        textColor=HexColor("#222222"),
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )
    quote_style = ParagraphStyle(
        "CustomQuote",
        parent=body_style,
        leftIndent=20,
        rightIndent=20,
        fontSize=10.5,
        textColor=HexColor("#666666"),
        fontName="ZHFont",
    )
    code_style = ParagraphStyle(
        "CustomCode",
        parent=body_style,
        fontName="ZHCodeFont",
        fontSize=10,
        leftIndent=10,
        backColor=HexColor("#F5F5F5"),
    )
    subtitle_style = ParagraphStyle(
        "CustomSub",
        parent=body_style,
        fontSize=9,
        textColor=HexColor("#999999"),
        alignment=TA_CENTER,
    )

    story = []
    if style == "report" and title:
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
        story.append(Spacer(1, 0.4 * cm))
    # clean 风格：不加标题和日期，文档从用户内容的第一行（通常是 # 姓名）开始

    blocks = _parse_markdown_blocks(markdown_text)
    for block in blocks:
        t = block["type"]
        if t == "h1":
            story.append(Paragraph(_pdf_inline(block["text"]), h1_style))
        elif t == "h2":
            story.append(Paragraph(_pdf_inline(block["text"]), h2_style))
        elif t == "h3":
            story.append(Paragraph(_pdf_inline(block["text"]), h3_style))
        elif t == "h4":
            story.append(Paragraph(_pdf_inline(block["text"]), h3_style))
        elif t == "p":
            story.append(Paragraph(_pdf_inline(block["text"]), body_style))
        elif t == "quote":
            story.append(Paragraph(_pdf_inline(block["text"]), quote_style))
            story.append(Spacer(1, 0.2 * cm))
        elif t == "list_ol":
            items = [Paragraph(_pdf_inline(it), body_style) for it in block["items"]]
            story.append(ListFlowable([ListItem(i) for i in items], bulletType="1"))
            story.append(Spacer(1, 0.2 * cm))
        elif t == "list_ul":
            items = [Paragraph(_pdf_inline(it), body_style) for it in block["items"]]
            story.append(ListFlowable([ListItem(i) for i in items], bulletType="bullet"))
            story.append(Spacer(1, 0.2 * cm))
        elif t == "code":
            for line in block["text"].split("\n"):
                story.append(Paragraph(_pdf_escape(line) or "&nbsp;", code_style))

    doc.build(story)
    return filepath


def _pdf_inline(text: str) -> str:
    """行内 Markdown 转 reportlab Paragraph 标签"""
    # 转义
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # **bold**
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # `code`
    text = re.sub(r"`(.+?)`", r'<font face="ZHCodeFont" color="#C03A3A">\1</font>', text)
    return text


def _pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_CHINESE_FONT_REGISTERED = False


def _register_chinese_font():
    """注册中文字体到 reportlab"""
    global _CHINESE_FONT_REGISTERED
    if _CHINESE_FONT_REGISTERED:
        return
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # 尝试常见的中文字体路径
        candidates = [
            ("/System/Library/Fonts/PingFang.ttc", "ZHFont"),
            ("/System/Library/Fonts/STHeiti Medium.ttc", "ZHFont"),
            ("/System/Library/Fonts/Hiragino Sans GB.ttc", "ZHFont"),
            ("/Library/Fonts/Songti.ttc", "ZHFont"),
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "ZHFont"),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "ZHFont"),
        ]
        # 找第一个存在的
        for path, name in candidates:
            if os.path.isfile(path):
                try:
                    pdfmetrics.registerFont(TTFont(name, path))
                    pdfmetrics.registerFont(TTFont("ZHCodeFont", path))
                    _CHINESE_FONT_REGISTERED = True
                    return
                except Exception:
                    continue
        # 如果都没找到，使用 reportlab 内置（可能中文乱码）
        _CHINESE_FONT_REGISTERED = True
    except Exception:
        pass

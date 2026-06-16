# 文档解析服务
# 统一处理 PDF / Word / TXT 上传文件，提取纯文本

import os
import uuid
from werkzeug.datastructures import FileStorage


def extract_text_from_upload(file: FileStorage) -> str:
    """
    从上传的 FileStorage 对象中提取纯文本
    支持：.pdf / .docx / .doc / .txt
    """
    if not file or not file.filename:
        raise ValueError("文件为空")

    filename = file.filename.lower()
    ext = os.path.splitext(filename)[1]

    # 临时保存文件以便解析
    tmp_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "static", "uploads")
    os.makedirs(tmp_dir, exist_ok=True)
    tmp_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex}_{file.filename}")
    file.save(tmp_path)

    try:
        if ext == ".pdf":
            return _extract_pdf(tmp_path)
        elif ext == ".docx":
            return _extract_docx(tmp_path)
        elif ext == ".doc":
            return _extract_doc_legacy(tmp_path)
        elif ext in (".txt", ".md"):
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            # 尝试按 txt 读取
            try:
                with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception:
                raise ValueError(f"不支持的文件格式：{ext}（支持 PDF/Word/TXT）")
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def _extract_pdf(path: str) -> str:
    """提取 PDF 文本"""
    text = ""
    # 优先 pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        return text
    except ImportError:
        pass
    except Exception:
        pass

    # 备选 PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except ImportError:
        raise ValueError("未安装 PDF 解析库（pdfplumber/PyPDF2）")
    except Exception as e:
        raise ValueError(f"PDF 解析失败：{str(e)}")


def _extract_docx(path: str) -> str:
    """提取 docx 文本"""
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        raise ValueError("未安装 python-docx（请 pip install python-docx）")
    except Exception as e:
        raise ValueError(f"Word 解析失败：{str(e)}")


def _extract_doc_legacy(path: str) -> str:
    """老 .doc 格式尽力解析（可能不完美）"""
    try:
        # 尝试用 python-docx
        return _extract_docx(path)
    except Exception:
        # 退化方案：按二进制读取可打印字符
        try:
            with open(path, "rb") as f:
                data = f.read()
            text = data.decode("utf-8", errors="ignore")
            # 简单清理
            return "".join(ch for ch in text if ch.isprintable() or ch in "\n\t")
        except Exception as e:
            raise ValueError(f"老 Word 格式解析失败：{str(e)}。建议另存为 .docx 或 .pdf")
